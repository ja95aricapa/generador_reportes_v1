import json
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Inches, Pt

# Función para convertir el HTML renderizado en un documento de Word
def html_to_docx(images_a, images_b, output_filename):
    doc = Document()

    # Añadir el contenido de la plantilla A
    doc.add_heading('Portfolio Overview', level=2)
    doc.add_paragraph("Este es un ejemplo de documento que contiene varios marcadores de posición para la imagen.")
    add_images_table(doc, images_a)

    # Añadir el contenido de la plantilla B
    doc.add_heading('Risk', level=2)
    doc.add_paragraph("Puedes proseguir leyendo este informe con texto dummy")
    add_images_table(doc, images_b)

    doc.save(output_filename)

def add_images_table(doc, images):
    row_count = len(images) // 2

    for i in range(row_count):
        row_images = images[i * 2: i * 2 + 2]
        table = doc.add_table(rows=1, cols=2)

        for row in table.rows:
            for cell in row.cells:
                for side in cell._element.xpath('.//w:tcBorders/*'):
                    side.set("w:val", "nil")

        for j, image in enumerate(row_images):
            cell = table.cell(0, j)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image["path"], width=Inches(3))
            cell.add_paragraph(image["id"])
        doc.add_paragraph()

def main():
    # Cargar el archivo JSON y la plantilla HTML
    with open("PRACTICE-REPOS/test_embajada/images/path.json", "r") as f:
        image_paths = json.load(f)

    env = Environment(loader=FileSystemLoader('.'))
    template_a = env.get_template('PRACTICE-REPOS/test_embajada/plantilla_prueba_v5_A.html')
    template_b = env.get_template('PRACTICE-REPOS/test_embajada/plantilla_prueba_v5_B.html')

    # Dividir las imágenes en dos listas para cada plantilla
    images = [
        {"id": key, "path": value}
        for key, value in image_paths.items()
    ]

    images_a = images[:4]
    images_b = images[4:]

    # Crear y guardar el documento de Word
    output_filename = "PRACTICE-REPOS/test_embajada/documento_con_imagenes_v6.docx"
    html_to_docx(images_a, images_b, output_filename)

# Entry point
if __name__ == "__main__":
    main()